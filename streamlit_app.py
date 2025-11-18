import streamlit as st
from openai import OpenAI

import streamlit as st
import pandas as pd
import os
from docx import Document  # from python-docx




import sqlite3
from glob import glob



import pdfplumber


DATA_DIR = "./federal_tax_documents/federal_forms"
DB_PATH = "./documents.db"


# ---------- DB SETUP ----------

def init_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT UNIQUE,
            file_path TEXT,
            file_type TEXT
        );
        """
    )

    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS extracted_rows (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            doc_id INTEGER,
            row_index INTEGER,
            col_index INTEGER,
            text TEXT,
            FOREIGN KEY(doc_id) REFERENCES documents(id)
        );
        """
    )

    conn.commit()
    conn.close()


def upsert_document(filename: str, file_path: str, file_type: str) -> int:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    cur.execute(
        """
        INSERT INTO documents (filename, file_path, file_type)
        VALUES (?, ?, ?)
        ON CONFLICT(filename) DO UPDATE SET file_path=excluded.file_path,
                                            file_type=excluded.file_type;
        """,
        (filename, file_path, file_type),
    )
    conn.commit()

    # get id
    cur.execute("SELECT id FROM documents WHERE filename = ?", (filename,))
    doc_id = cur.fetchone()[0]

    conn.close()
    return doc_id


def clear_extracted_for_doc(doc_id: int):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("DELETE FROM extracted_rows WHERE doc_id = ?", (doc_id,))
    conn.commit()
    conn.close()


def insert_extracted_rows(doc_id: int, rows: list[dict]):
    """
    rows: list of dicts {row_index, col_index, text}
    """
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.executemany(
        """
        INSERT INTO extracted_rows (doc_id, row_index, col_index, text)
        VALUES (?, ?, ?, ?)
        """,
        [(doc_id, r["row_index"], r["col_index"], r["text"]) for r in rows],
    )
    conn.commit()
    conn.close()


# ---------- EXTRACTION HELPERS ----------

def extract_pdf_tables(file_path: str) -> list[dict]:
    """
    Extract all tables from a PDF as list of {row_index, col_index, text}.
    """
    extracted = []
    row_counter = 0

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            if not tables:
                continue
            for table in tables:
                for row in table:
                    if row is None:
                        continue
                    for col_idx, cell in enumerate(row):
                        text = (cell or "").strip()
                        extracted.append(
                            {
                                "row_index": row_counter,
                                "col_index": col_idx,
                                "text": text,
                            }
                        )
                    row_counter += 1

    return extracted


def extract_docx_lines(file_path: str) -> list[dict]:
    """
    Extract paragraphs from a DOCX as a single 'column' table:
    col_0 = text
    """
    extracted = []
    doc = Document(file_path)

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        extracted.append(
            {
                "row_index": idx,
                "col_index": 0,
                "text": text,
            }
        )

    return extracted


# ---------- PROCESS ALL DOCUMENTS ----------

def process_all_documents():
    """
    Scan DATA_DIR for .pdf and .docx,
    extract content, and store in SQLite.
    """
    pdf_files = glob(os.path.join(DATA_DIR, "*.pdf"))
    docx_files = glob(os.path.join(DATA_DIR, "*.docx"))

    total_docs = 0
    total_rows = 0

    for file_path in pdf_files + docx_files:
        filename = os.path.basename(file_path)
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            file_type = "pdf"
            rows = extract_pdf_tables(file_path)
        elif ext == ".docx":
            file_type = "docx"
            rows = extract_docx_lines(file_path)
        else:
            continue  # skip unknown

        doc_id = upsert_document(filename, file_path, file_type)
        clear_extracted_for_doc(doc_id)

        if rows:
            insert_extracted_rows(doc_id, rows)
            total_rows += len(rows)

        total_docs += 1

    return total_docs, total_rows


def get_documents_df() -> pd.DataFrame:
    conn = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query("SELECT id, filename, file_type FROM documents", conn)
    conn.close()
    return df


def get_extracted_for_doc(doc_id: int) -> pd.DataFrame:
    conn = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query(
        """
        SELECT row_index, col_index, text
        FROM extracted_rows
        WHERE doc_id = ?
        ORDER BY row_index, col_index
        """,
        conn,
        params=(doc_id,),
    )
    conn.close()
    return df


# ---------- STREAMLIT APP ----------



# init DB on startup
init_db()

#st.sidebar.header("Actions")


docs, rows = process_all_documents()
st.success(f"Processed {docs} document(s), extracted {rows} cell(s).")

# List all documents from DB
if os.path.exists(DB_PATH):
    docs_df = get_documents_df()
else:
    docs_df = pd.DataFrame(columns=["id", "filename", "file_type"])

st.subheader("Documents in persistent DB")

if docs_df.empty:
    st.info("No documents indexed yet.")
else:
    st.dataframe(docs_df)

    # Select a document to view
    doc_options = {f'{row["filename"]} ({row["file_type"]})': row["id"] for _, row in docs_df.iterrows()}
    selected_label = st.selectbox("Select a document to view extracted data:", list(doc_options.keys()))
    selected_doc_id = doc_options[selected_label]

    # Show extracted rows for that document
    extracted_df = get_extracted_for_doc(selected_doc_id)

    st.subheader(f"Extracted table for: {selected_label}")
    if extracted_df.empty:
        st.warning("No extracted data for this document.")
    else:
        st.dataframe(extracted_df)

    # "Link" / download button for the document
    # (This is the closest thing to a link for local files in Streamlit Cloud)
    # We look up the file_path from DB:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT file_path FROM documents WHERE id = ?", (selected_doc_id,))
    file_path = cur.fetchone()[0]
    conn.close()

    if os.path.exists(file_path):
        with open(file_path, "rb") as f:
            st.download_button(
                label="Download original document",
                data=f,
                file_name=os.path.basename(file_path),
                mime="application/octet-stream",
            )
    else:
        st.error(f"File not found on disk: {file_path}")



# Optional: install if you want PDF/DOCX text extraction
# pip install pypdf2 python-docx
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

st.set_page_config(page_title="AI Tax Assistant", page_icon="💬")

st.title("💬 AI Assistant Chatbot")
st.write(
    "This is a simple chatbot that uses OpenAI's GPT-3.5 model to help fill your tax forms."
)

# ---------------------------------------------------------
# 🔹 Upload Area: 2 Columns (Documents & Images)
# ---------------------------------------------------------
col1, col2 = st.columns(2)

with col1:
    st.subheader("📄 Upload a document")
    doc_file = st.file_uploader(
        "Upload a PDF / TXT / DOCX",
        type=["pdf", "txt", "docx"],
        key="doc_uploader",
    )

with col2:
    st.subheader("🖼️ Upload an image")
    img_file = st.file_uploader(
        "Upload an image",
        type=["png", "jpg", "jpeg"],
        key="img_uploader",
    )

uploaded_text = None  # text extracted from uploaded document (if any)

# ---------------------------------------------------------
# 📄 Handle document upload + preview
# ---------------------------------------------------------
if doc_file is not None:
    st.markdown(f"**Document uploaded:** `{doc_file.name}`")

    if doc_file.type == "text/plain":
        # Simple text file
        uploaded_text = doc_file.read().decode("utf-8", errors="ignore")

    elif doc_file.type == "application/pdf":
        if PdfReader is None:
            st.warning("PyPDF2 is not installed. Run `pip install pypdf2` to read PDFs.")
        else:
            reader = PdfReader(doc_file)
            text_chunks = []
            for page in reader.pages[:5]:  # limit pages to avoid huge output
                text_chunks.append(page.extract_text() or "")
            uploaded_text = "\n".join(text_chunks)

    elif doc_file.type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]:
        if Document is None:
            st.warning("python-docx is not installed. Run `pip install python-docx`.")
        else:
            doc = Document(doc_file)
            paragraphs = [p.text for p in doc.paragraphs]
            uploaded_text = "\n".join(paragraphs)

    # Show the extracted / raw text (if any)
    if uploaded_text:
        with st.expander("📄 Preview extracted document text"):
            st.text(uploaded_text)
    else:
        st.info("Document uploaded but no text could be extracted.")

# ---------------------------------------------------------
# 🖼️ Handle image upload + preview
# ---------------------------------------------------------
if img_file is not None:
    st.markdown(f"**Image uploaded:** `{img_file.name}`")
    st.image(img_file, caption="Uploaded image", use_container_width=True)

# ---------------------------------------------------------
# 🔑 OpenAI setup
# ---------------------------------------------------------
openai_api_key = os.getenv("OPENAI_API_KEY")

if not openai_api_key:
    st.info("Update your key in the environment variable OPENAI_API_KEY.", icon="🗝️")
else:
    # Create an OpenAI client.
    client = OpenAI(api_key=openai_api_key)

    # Session state for chat messages
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Display previous messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Chat input
    prompt = st.chat_input("Do you want to file your taxes?")
    if prompt:

        # Optionally inject uploaded document text into the context
        if uploaded_text:
            prompt_with_context = (
                "The user has uploaded the following document text:\n\n"
                f"{uploaded_text}\n\n"
                f"User question: {prompt}"
            )
        else:
            prompt_with_context = prompt

        # Store and display user's message
        st.session_state.messages.append(
            {"role": "user", "content": prompt_with_context}
        )
        with st.chat_message("user"):
            st.markdown(prompt)

        # Generate response
        stream = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": m["role"], "content": m["content"]}
                for m in st.session_state.messages
            ],
            stream=True,
        )

        # Stream response and save it
        with st.chat_message("assistant"):
            response = st.write_stream(stream)
        st.session_state.messages.append(
            {"role": "assistant", "content": response}
        )
